
from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
font.color.rgb = RGBColor(0, 0, 0)

for i in range(1, 5):
    h_style = doc.styles.get(f'Heading {i}')
    if h_style:
        h_style.font.name = 'Times New Roman'
        h_style.font.size = Pt(14)
        h_style.font.color.rgb = RGBColor(0, 0, 0)
        h_style.font.bold = True

def add_p(text, bold_prefix=None, is_bullet=False):
    p = doc.add_paragraph(style='List Bullet' if is_bullet else 'Normal')
    p.paragraph_format.line_spacing = 1.5
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0, 0, 0)
    
    r2 = p.add_run(text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0, 0, 0)

def add_h(text, level=2):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.line_spacing = 1.5

add_h('CropNGo Documentation', level=1)
add_h('Table of Contents', level=2)
add_p('1. Introduction
2. The Solution
3. Technology Stack
4. Future Growth')

add_h('1. Introduction', level=2)
add_p('Agriculture is the backbone of human civilization, but it faces unprecedented challenges today. Climate change, supply chain disruptions, and resource scarcity threaten global food security. In Malaysia and around the world, adopting Agritech (Agricultural Technology) is no longer just an option; it is a necessity. Agritech empowers farmers to optimize yields, reduce waste, and manage resources efficiently. Despite these advancements, many traditional farmers remain disconnected from modern markets, struggling to sell their produce directly to consumers or collaborate with other stakeholders. CropNGo bridges this gap by providing a centralized platform that digitalizes agriculture, ensuring sustainability and food security for the future.')

add_h('2. The Solution', level=2)
add_p('CropNGo serves as the ultimate digital ecosystem to solve the unresolved issues in modern agriculture. The platform is designed to connect farmers, vendors, and consumers, fostering a thriving community. Key features include:')

add_p('CropNGo integrates advanced AI to assist users with agricultural queries, provide smart recommendations, and facilitate intelligent searches for crops and farming resources. The AI acts as a virtual farming assistant, making complex data accessible to everyone.', bold_prefix='AI Features: ', is_bullet=True)
add_p('Agriculture thrives on community knowledge. The event page allows users to discover and participate in agricultural events, workshops, and seminars, promoting continuous learning and networking.', bold_prefix='Event Page: ', is_bullet=True)
add_p('A dedicated marketplace enables farmers to sell their fresh produce and farming equipment directly to buyers. This eliminates the middleman, ensuring fair prices for farmers and fresh goods for consumers.', bold_prefix='Shop Page: ', is_bullet=True)
add_p('When users cannot find specific items, they can post requests on the wanted listing. This demand-driven feature helps suppliers understand market needs and fulfills specific consumer demands quickly.', bold_prefix='Wanted Listing: ', is_bullet=True)
add_p('At its core, CropNGo is about community. By connecting farmers, vendors, and consumers seamlessly, the platform builds trust, encourages collaboration, and creates a resilient agricultural supply chain.', bold_prefix='Emphasizing Connection: ', is_bullet=True)

add_h('3. Technology Stack', level=2)
add_p('To deliver a robust, scalable, and user-friendly experience, CropNGo utilizes a modern technology stack for both its frontend and backend architectures:')

add_p('The user interface is built using React with TypeScript. This combination ensures a highly responsive, type-safe, and dynamic web experience that works flawlessly across different devices.', bold_prefix='Frontend: ', is_bullet=True)
add_p('Firebase is used for real-time database management, user authentication, and secure cloud storage. Its scalable infrastructure ensures that user data and marketplace transactions are handled securely and efficiently.', bold_prefix='Backend & Database: ', is_bullet=True)
add_p('The platform leverages the Google Gemini API to power its intelligent chat, automated data processing, and smart search features. This brings cutting-edge generative AI capabilities directly to the users.', bold_prefix='AI Integration: ', is_bullet=True)
add_p('The Google Maps API is integrated to provide accurate geolocation services, enabling users to find nearby events, farmers, and delivery routes easily.', bold_prefix='Location Services: ', is_bullet=True)

add_h('4. Future Growth', level=2)
add_p('CropNGo is continuously evolving to meet the growing needs of the agricultural community. Our vision for future growth includes:')

add_p('We plan to introduce a system that allows users to provide financial incentives and donations to poor farmers, offering them the vital support needed to sustain their livelihoods.', bold_prefix='Incentives and Donations: ', is_bullet=True)
add_p('CropNGo will organize huge agricultural events and expos, bringing together industry experts, technology providers, and farmers to showcase the latest innovations and foster large-scale networking.', bold_prefix='Agricultural Events: ', is_bullet=True)
add_p('To empower the next generation of farmers, we will integrate farming courses, interactive quizzes, and official certificates. This educational hub will promote modern farming techniques and sustainability.', bold_prefix='Education and Certification: ', is_bullet=True)
add_p('To ensure trust and reliability within the marketplace, an enhanced rating system will be implemented. This system will use advanced algorithms to detect and prevent ethical violations, such as spamming or fake reviews, ensuring a fair environment for all users.', bold_prefix='Enhanced Rating System: ', is_bullet=True)
add_p('Future updates will include integration with Internet of Things (IoT) sensors, allowing farmers to monitor soil moisture, temperature, and crop health directly from the CropNGo app.', bold_prefix='IoT Integration for Smart Farming: ', is_bullet=True)
add_p('By analyzing historical data and weather patterns, the platform will offer predictive analytics to help farmers forecast their crop yields and optimize harvest times.', bold_prefix='Predictive Yield Analytics: ', is_bullet=True)
add_p('To streamline the supply chain further, we will partner with local logistics companies to offer integrated, affordable, and trackable delivery options for all marketplace transactions.', bold_prefix='Logistics and Delivery Partnership: ', is_bullet=True)

doc.save('CropNGo_Documentation.docx')
